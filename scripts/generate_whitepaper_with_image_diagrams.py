from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


def _root() -> Path:
    return Path(__file__).resolve().parents[1]


def _font(size: int) -> ImageFont.ImageFont:
    try:
        return ImageFont.truetype("arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def _draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str) -> None:
    draw.rounded_rectangle(xy, radius=12, outline=(30, 30, 30), width=3, fill=(245, 248, 255))
    font = _font(20)
    x1, y1, x2, y2 = xy
    max_width = x2 - x1 - 30
    words = text.split()
    lines: list[str] = []
    cur = ""
    for w in words:
        test = (cur + " " + w).strip()
        if draw.textlength(test, font=font) <= max_width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    total_h = len(lines) * 28
    y = y1 + ((y2 - y1 - total_h) // 2)
    for line in lines:
        tw = int(draw.textlength(line, font=font))
        draw.text((x1 + ((x2 - x1 - tw) // 2), y), line, fill=(20, 20, 20), font=font)
        y += 28


def _arrow(draw: ImageDraw.ImageDraw, x1: int, y1: int, x2: int, y2: int) -> None:
    draw.line((x1, y1, x2, y2), fill=(50, 70, 180), width=4)
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    head = 14
    a1 = angle + 2.6
    a2 = angle - 2.6
    p1 = (x2 + head * math.cos(a1), y2 + head * math.sin(a1))
    p2 = (x2 + head * math.cos(a2), y2 + head * math.sin(a2))
    draw.polygon([(x2, y2), p1, p2], fill=(50, 70, 180))


def _diagram_hybrid(path: Path) -> None:
    img = Image.new("RGB", (2000, 1200), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_font = _font(34)
    d.text((40, 24), "Dual-Chain Hybrid Architecture", fill=(0, 0, 0), font=title_font)

    boxes = {
        "user": (60, 120, 460, 260),
        "adapter": (520, 120, 940, 260),
        "olas": (1020, 120, 1500, 260),
        "norm": (520, 360, 940, 500),
        "celo": (520, 600, 940, 760),
        "settle": (1020, 600, 1500, 760),
        "trace": (60, 860, 1500, 1060),
    }
    _draw_box(d, boxes["user"], "Operator / API Request")
    _draw_box(d, boxes["adapter"], "Public Adapter + Boundary Logic")
    _draw_box(d, boxes["olas"], "Gnosis Olas Mech Marketplace")
    _draw_box(d, boxes["norm"], "Task Normalization + Correlation")
    _draw_box(d, boxes["celo"], "Celo Private Marketplace")
    _draw_box(d, boxes["settle"], "Settlement + Withdrawals")
    _draw_box(d, boxes["trace"], "Communication Trace + Reports + Proof Artifacts")

    _arrow(d, 460, 190, 520, 190)
    _arrow(d, 940, 190, 1020, 190)
    _arrow(d, 730, 260, 730, 360)
    _arrow(d, 730, 500, 730, 600)
    _arrow(d, 940, 680, 1020, 680)
    _arrow(d, 730, 760, 730, 860)
    _arrow(d, 1260, 760, 1260, 860)
    _arrow(d, 1260, 260, 1260, 860)

    img.save(path)


def _diagram_settlement(path: Path) -> None:
    img = Image.new("RGB", (2000, 1100), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_font = _font(34)
    d.text((40, 24), "Private Celo Settlement Lifecycle", fill=(0, 0, 0), font=title_font)

    steps = [
        ("Requester creates task (escrow)", (120, 160, 560, 280)),
        ("Worker accepts and submits result", (700, 160, 1200, 280)),
        ("Validator scores and finalizes", (1340, 160, 1840, 280)),
        ("Protocol + finance fees allocated", (320, 440, 900, 580)),
        ("Worker payout + requester refund", (1080, 440, 1680, 580)),
        ("Withdrawals by entitled addresses", (560, 760, 1440, 930)),
    ]
    for text, box in steps:
        _draw_box(d, box, text)

    _arrow(d, 560, 220, 700, 220)
    _arrow(d, 1200, 220, 1340, 220)
    _arrow(d, 1540, 280, 1380, 440)
    _arrow(d, 1540, 280, 760, 440)
    _arrow(d, 900, 510, 1080, 510)
    _arrow(d, 610, 580, 860, 760)
    _arrow(d, 1380, 580, 1140, 760)

    img.save(path)


def _append_diagrams_to_docx(source: Path, out: Path, d1: Path, d2: Path) -> None:
    doc = Document(str(source))
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("Architecture Diagrams (Image Rendered)")
    r.bold = True
    r.font.size = Pt(18)

    doc.add_paragraph("Figure 1. Dual-chain hybrid architecture")
    doc.add_picture(str(d1), width=Inches(7.2))
    doc.add_paragraph("Figure 2. Private Celo settlement lifecycle")
    doc.add_picture(str(d2), width=Inches(7.2))
    doc.save(str(out))


def main() -> None:
    root = _root()
    base_doc = root / "AgenticSwarmMarketplace_Whitepaper_Expanded.docx"
    out_doc = root / "AgenticSwarmMarketplace_Whitepaper_Expanded_Diagrams.docx"
    diag_dir = root / "architecture_diagrams"
    diag_dir.mkdir(exist_ok=True)

    d1 = diag_dir / "dual_chain_hybrid_architecture.png"
    d2 = diag_dir / "private_settlement_lifecycle.png"
    _diagram_hybrid(d1)
    _diagram_settlement(d2)
    _append_diagrams_to_docx(base_doc, out_doc, d1, d2)
    print(f"Wrote {d1}")
    print(f"Wrote {d2}")
    print(f"Wrote {out_doc}")


if __name__ == "__main__":
    main()
