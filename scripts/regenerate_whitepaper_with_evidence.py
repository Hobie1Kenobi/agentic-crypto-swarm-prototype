from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _root() -> Path:
    return Path(__file__).resolve().parents[1]


def _report_json(root: Path, name: str) -> Path:
    ar = root / "artifacts" / "reports" / name
    if ar.exists():
        return ar
    return root / name


def _comm_trace_json(root: Path) -> Path:
    p = root / "artifacts" / "communication" / "communication_trace.json"
    if p.exists():
        return p
    return root / "communication_trace.json"


def _add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.rstrip() + "\n")
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level <= 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(f"- {text}")


def _safe(v: Any, fallback: str = "n/a") -> str:
    if v is None or v == "":
        return fallback
    return str(v)


def main() -> None:
    root = _root()
    source_docx = Path(r"C:\Users\hobie\Downloads\AgenticSwarmMarketplace_Whitepaper.docx")
    out_docx = root / "AgenticSwarmMarketplace_Whitepaper_Expanded.docx"

    preflight = _read_json(root / "olas_preflight_report.json")
    live = _read_json(root / "olas_live_attempt_report.json")
    hybrid = _read_json(root / "hybrid_gnosis_celo_report.json")
    private = _read_json(root / "celo_sepolia_task_market_report.json")
    trace = _read_json(root / "communication_trace.json")
    public = _read_json(root / "public_adapter_run_report.json")

    if source_docx.exists():
        doc = Document(str(source_docx))
    else:
        doc = Document()
        _add_heading(doc, "Agentic Swarm Marketplace Whitepaper", level=1)
        doc.add_paragraph("Base source file not found; generated from project artifacts.")

    doc.add_page_break()
    _add_heading(doc, "Expanded Implementation Evidence", level=1)
    doc.add_paragraph(
        "This section extends the original whitepaper with implementation proof drawn from live project artifacts. "
        "It documents what has been deployed, validated, and executed across private Celo settlement and public Olas integration."
    )
    doc.add_paragraph(f"Generated at: {datetime.now(timezone.utc).isoformat()} (UTC)")

    _add_heading(doc, "Proof-of-Work Summary", level=2)
    bullets = [
        f"Private chain settlement status: {_safe(((private.get('task') or {}).get('task_status') or {}).get('name'))}",
        f"Private internal task id: {_safe((private.get('task') or {}).get('task_id'))}",
        f"External Olas boundary: {_safe((live.get('result') or {}).get('boundary') or live.get('boundary'))}",
        f"Olas preflight passed: {_safe(preflight.get('ok'))}",
        f"Hybrid settlement consistency: {_safe((hybrid.get('settlement') or {}).get('settlement_matches_expected'))}",
    ]
    for b in bullets:
        _add_bullet(doc, b)

    _add_heading(doc, "Evidence by Artifact", level=2)
    artifact_rows = [
        ("olas_preflight_report.json", "Environment and funding readiness gate"),
        ("olas_live_attempt_report.json", "One-shot live Olas request attempt and blocker capture"),
        ("public_adapter_run_report.json", "Public intake adapter execution output"),
        ("communication_trace.json", "Cross-boundary event timeline and correlation"),
        ("celo_sepolia_task_market_report.json", "Private Celo task lifecycle and settlement tx hashes"),
        ("hybrid_gnosis_celo_report.json", "Hybrid-mode synthesis: external boundary + internal settlement"),
        ("private_celo_validation_report.md", "Human-readable private-chain validation narrative"),
    ]
    for name, desc in artifact_rows:
        p = doc.add_paragraph("- ")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    _add_heading(doc, "Architecture Diagram: Dual-Chain Hybrid Mode", level=2)
    _add_code_block(
        doc,
        """
+----------------------+         +------------------------------+
| Operator / API User  |         | .env + Activation Workflow   |
| submits prompt/task  |         | (mechx, EOA, funding checks) |
+----------+-----------+         +---------------+--------------+
           |                                     |
           v                                     v
+----------------------+                +----------------------+
| Public Adapter Layer |                | Olas Preflight Gate  |
| normalize + boundary |                | config + balance     |
+----------+-----------+                +-----------+----------+
           |                                        |
           |  real_external_integration OR          |
           |  mocked_external_replay                |
           v                                        |
+-------------------------------+                   |
| Gnosis / Olas Mech Marketplace|<------------------+
| request id + tx hash (if live)|
+---------------+---------------+
                |
                v
+--------------------------------------------+
| Internal Task Normalization + Correlation  |
| maps external request -> internal task id  |
+----------------+---------------------------+
                 |
                 v
+--------------------------------------------+
| Celo Private Marketplace (contract level)  |
| create -> accept -> submit -> score -> fin |
+----------------+---------------------------+
                 |
                 v
+--------------------------------------------+
| Settlement Accounting + Withdrawals         |
| protocol_fee / finance_fee / payout/refund |
+--------------------------------------------+
""",
    )

    _add_heading(doc, "Architecture Diagram: Settlement Flow", level=2)
    _add_code_block(
        doc,
        """
Requester (ROOT_STRATEGIST)
   -> createTask (escrow)
Worker (DEPLOYER)
   -> acceptTask
   -> submitResult
Validator (FINANCE_DISTRIBUTOR)
   -> submitTaskScore
   -> finalizeTask
Entitlements created:
   protocol_fee      -> treasury
   finance_fee       -> finance_distributor
   worker_payout     -> worker
   requester_refund  -> requester
Each address -> withdraw()
""",
    )

    _add_heading(doc, "Execution Evidence and Correlation", level=2)
    corr = hybrid.get("correlation") or {}
    by_cat = ((hybrid.get("settlement") or {}).get("by_category") or {})
    doc.add_paragraph(f"External request id: {_safe(corr.get('external_request_id'))}")
    doc.add_paragraph(f"External tx hash: {_safe(corr.get('external_tx_hash'))}")
    doc.add_paragraph(f"Internal Celo task id: {_safe(corr.get('internal_task_id'))}")
    doc.add_paragraph(f"Public chain id: {_safe((trace.get('public_chain_id')))}")
    doc.add_paragraph(f"Private chain id: {_safe((trace.get('private_chain_id')))}")
    for category in ("protocol_fee", "finance_fee", "worker_payout", "requester_refund"):
        c = by_cat.get(category) or {}
        doc.add_paragraph(
            f"{category}: expected={_safe(c.get('expected_wei'))} wei, "
            f"actual_pending={_safe(c.get('actual_pending_wei'))} wei",
        )

    txs = corr.get("internal_tx_hashes") or []
    if txs:
        _add_heading(doc, "Internal Transaction Hashes", level=3)
        for t in txs:
            _add_bullet(doc, f"{_safe(t.get('name'))} ({_safe(t.get('role'))}): {_safe(t.get('tx_hash'))}")

    _add_heading(doc, "Live Integration Status", level=2)
    live_ok = bool(live.get("ok"))
    live_err = _safe((live.get("result") or {}).get("error"))
    doc.add_paragraph(f"One-shot live Olas request success: {live_ok}")
    doc.add_paragraph(f"Live boundary: {_safe(live.get('boundary'))}")
    doc.add_paragraph(f"Last live attempt blocker/error: {live_err}")
    doc.add_paragraph(
        "The system preserves deterministic fallback behavior: if live external execution is blocked, "
        "the hybrid flow remains operational through mocked_external_replay while private Celo settlement stays real."
    )

    _add_heading(doc, "Risk Register and Next Milestones", level=2)
    risks = [
        "External request funding sensitivity on Gnosis (gas + marketplace fee + tip behavior).",
        "Tool-to-mech compatibility must be validated before request send.",
        "Address-role segregation requires dedicated keys and deploy-time treasury alignment.",
    ]
    for r in risks:
        _add_bullet(doc, r)
    doc.add_paragraph("Next milestone: complete one successful real_external_integration request and regenerate hybrid live report.")

    _add_heading(doc, "Appendix: Artifact Index", level=2)
    reports_dir = root / "artifacts" / "reports"
    if reports_dir.is_dir():
        for p in sorted(reports_dir.glob("*report*.md")):
            _add_bullet(doc, str(p.relative_to(root)))
        for p in sorted(reports_dir.glob("*report*.json")):
            _add_bullet(doc, str(p.relative_to(root)))
    else:
        for p in sorted(root.glob("*report*.md")):
            _add_bullet(doc, p.name)
        for p in sorted(root.glob("*report*.json")):
            _add_bullet(doc, p.name)
    comm_dir = root / "artifacts" / "communication"
    for name in ["communication_trace.md", "communication_trace.json"]:
        cp = comm_dir / name if comm_dir.is_dir() else root / name
        if cp.exists():
            _add_bullet(doc, str(cp.relative_to(root)))
    for name in ["olas_activation_checklist.md", "olas_funding_address.txt"]:
        if (root / name).exists():
            _add_bullet(doc, name)

    doc.save(str(out_docx))
    print(f"Wrote {out_docx}")


if __name__ == "__main__":
    main()
